from docx import Document
from pathlib import Path
from config import data, TEMPLATES_DIR, OUTPUT_DIR
from utils import format_filename

def replace_all_in_paragraph(par):
    full = ''.join(r.text for r in par.runs)
    return full

def set_paragraph_text(par, text):
    for r in par.runs:
        r.text = ''
    if par.runs:
        par.runs[0].text = text
        
        for r in par.runs[1:]:
            r.text = ''
    else:
        par.add_run(text)

def replace_text(doc, data: dict) -> str:

    for par in doc.paragraphs:
        orig = replace_all_in_paragraph(par)
        new = orig
        for key, value in data.items():
            if key in new:
                new = new.replace(key, value)
        if new != orig:
            set_paragraph_text(par, new)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    orig = replace_all_in_paragraph(par)
                    new = orig
                    for key, value in data.items():
                        if key in new:
                            new = new.replace(key, value)
                    if new != orig:
                        set_paragraph_text(par, new)


def fill_template(template_path, output_path, data):
    tp = Path(template_path)
    op = Path(output_path)
    if not tp.exists():
        raise FileNotFoundError(f"Template não encontrado: {tp}")
    op.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(str(tp))
    replace_text(doc, data)
    doc.save(str(op))

def generate_document(user_name: str, passwd: str, process_num: str) -> str:

    data["{NAME}"] = user_name
    data["{PASSWD}"] = passwd
    data["{PROCESS}"] = process_num

    output_file = OUTPUT_DIR / f"{format_filename(user_name)}.docx"

    fill_template(TEMPLATES_DIR, output_file, data)
    
    return str(output_file)




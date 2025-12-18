from pathlib import Path
from docx import Document


def _get_full_paragraph_text(par):
    return "".join(run.text for run in par.runs)


def _set_paragraph_text(par, text):
    for run in par.runs:
        run.text = ""
    if par.runs:
        par.runs[0].text = text
    else:
        par.add_run(text)


def _replace_text(doc, data: dict):
    for par in doc.paragraphs:
        original = _get_full_paragraph_text(par)
        new = original

        for key, value in data.items():
            if key in new:
                new = new.replace(key, value)

        if new != original:
            _set_paragraph_text(par, new)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    original = _get_full_paragraph_text(par)
                    new = original

                    for key, value in data.items():
                        if key in new:
                            new = new.replace(key, value)

                    if new != original:
                        _set_paragraph_text(par, new)


def fill_template(template_path: Path, output_path: Path, data: dict):
    if not template_path.exists():
        raise FileNotFoundError(f"Template não encontrado: {template_path}")

    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(template_path))
    _replace_text(doc, data)
    doc.save(str(output_path))

from tkinter import ttk, messagebox

from docx import Document
from pathlib import Path
from common.config import data, TEMPLATES_DIR, OUTPUT_DIR
from common.utils import format_filename
from common.generator import generate_password

def replace_all_in_paragraph(par):
    full = ''.join(r.text for r in par.runs)
    return full

def set_paragraph_text(par, text):
    for r in par.runs:
        r.text = ''
    if par.runs:
        par.runs[0].text = text
        for r in par.runs[1:]:
            r.text = ''
    else:
        par.add_run(text)

def replace_text(doc, data: dict) -> str:
    for par in doc.paragraphs:
        orig = replace_all_in_paragraph(par)
        new = orig
        for key, value in data.items():
            if key in new:
                new = new.replace(key, value)
        if new != orig:
            set_paragraph_text(par, new)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    orig = replace_all_in_paragraph(par)
                    new = orig
                    for key, value in data.items():
                        if key in new:
                            new = new.replace(key, value)
                    if new != orig:
                        set_paragraph_text(par, new)

def fill_template(template_path, output_path, data):
    tp = Path(template_path)
    op = Path(output_path)

    if not tp.exists():
        raise FileNotFoundError(f"Template não encontrado: {tp}")

    op.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(tp))
    replace_text(doc, data)
    doc.save(str(op))

def generate_document(user_name: str, passwd: str, process_num: str) -> str:
    data["{NAME}"] = user_name
    data["{PASSWD}"] = passwd
    data["{PROCESS}"] = process_num

    output_file = OUTPUT_DIR / f"{format_filename(user_name)}.docx"

    fill_template(TEMPLATES_DIR, output_file, data)

    return str(output_file)

def tab_create_docs(app, container):
    frame = ttk.Frame(container, padding=40)
    frame.place(relx=0.5, rely=0.5, anchor="center")

    ttk.Label(
        frame,
        text="Gerador de Relatórios",
        font=("Times New Roman", 16, "bold"),
        foreground="#F5F5F5"
    ).pack(pady=(0, 25))

    ttk.Label(frame, text="Nome (formato: nome.sobrenome)").pack(pady=(10, 5))
    name_entry = ttk.Entry(
        frame,
        width=35,
        font=("Times New Roman", 10)
    )
    name_entry.pack(ipady=6)

    ttk.Label(frame, text="Número do processo (formato: XXXXXX)").pack(pady=(20, 5))
    process_entry = ttk.Entry(
        frame,
        width=35,
        font=("Times New Roman", 10)
    )
    process_entry.pack(ipady=6)

    def on_generate_document():
        try:
            user_name = name_entry.get().strip()
            passwd = generate_password(user_name)
            path = generate_document(user_name, passwd, process_entry.get())

            messagebox.showinfo(
                "Sucesso",
                f"Relatório gerado com sucesso!\n\nCaminho:\n{path}"
            )
        except Exception as e:
            messagebox.showerror("Erro", str(e))

    ttk.Button(
        frame,
        text="Gerar Relatório",
        command=on_generate_document,
        style="Accent.TButton"
    ).pack(pady=30)

    ttk.Separator(frame, orient="horizontal").pack(fill="x", pady=15)
    ttk.Label(
        frame,
        text="RelSyS © 2025",
        font=("Times New Roman", 9, "italic"),
        foreground=app.colors["subtext_color"]
    ).pack()

from docx import Document


def replace_text(doc, data):
    for element in doc.paragraphs + [
        p for t in doc.tables for r in t.rows for c in r.cells for p in c.paragraphs
    ]:
        for key, value in data.items():
            if key in element.text:
                for run in element.runs:
                    run.text = run.text.replace(key, value)


def fill_template(template_path, output_path, data):
    doc = Document(template_path)
    replace_text(doc, data)
    doc.save(output_path)


data = {
    "{{LOGIN_BASE}}": "eduardo.s",
    "{{EMAIL}}": "eduardo.s@workongroup.com.br",
    "{{PASSWORD_WINDOWS}}": "NoS@work2025",
    "{{PASSWORD_EMAIL}}": "Mudar@2025",
    "{{PASSWORD_GINFOR}}": "1234",
    "{{PASSWORD_TEAMS}}": "Mudar@2025",
    "{{PASSWORD_VBD}}": "NoS@work2025",
    "{{PASSWORD_FORTCLIENT}}": "workon2025",
    "{{PROCESS}}": "12345",
    "{{ASSET}}": "MAQ-789",
}

fill_template(
    "Bem-vindo a Work On Atualizado.docx", "/path/to/save/Final_Document.docx", data
)

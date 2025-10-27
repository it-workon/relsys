from docx import Document
from pathlib import Path

def replace_all_in_paragraph(par):
    full = ''.join(r.text for r in par.runs)
    return full

def set_paragraph_text(par, text):
    for r in par.runs:
        r.text = ''
    if par.runs:
        par.runs[0].text = text
        
        for r in par.runs[1:]:
            r.text = ''
    else:
        par.add_run(text)

def replace_text(doc, data):

    for par in doc.paragraphs:
        orig = replace_all_in_paragraph(par)
        new = orig
        for key, value in data.items():
            if key in new:
                new = new.replace(key, value)
        if new != orig:
            set_paragraph_text(par, new)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    orig = replace_all_in_paragraph(par)
                    new = orig
                    for key, value in data.items():
                        if key in new:
                            new = new.replace(key, value)
                    if new != orig:
                        set_paragraph_text(par, new)

def fill_template(template_path, output_path, data):
    tp = Path(template_path)
    op = Path(output_path)
    if not tp.exists():
        raise FileNotFoundError(f"Template não encontrado: {tp}")
    op.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(str(tp))
    replace_text(doc, data)
    doc.save(str(op))


data = {
    "{{USER_NAME}}": "fabia.martins",

    "{{DOMAIN_ORG}}": "workongroup.com.br",

    "{{USER_PASSWD}}": "FaM@work2025",

    "{{EMAIL_PASSWD}}": "Mudar@2025",

    "{{GI_PASSWD}}": "1234",

    "{{TEAMS_PASSWD}}": "Mudar@2025",

    "{{VBD_PASSWD}}": "FaM@work2025",

    "{{FORTICLIENT_PASSWD}}": "workon2025",

    "{{PROCESS_NUMBER}}": "12345",

    "{{ASSET}}": "EMC-126841",
}

fill_template(
    r"C:\Users\eduardo.andrade\Documents\Automator\relsys\environment\welcome-model.docx",
    r"C:\Users\eduardo.andrade\Documents\fabia-martins.docx",
    data
)

